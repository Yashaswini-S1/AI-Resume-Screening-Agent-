import os
import fitz  # PyMuPDF
import docx

def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"Error parsing PDF {file_path}: {e}")
        raise e
    return text

def parse_docx(file_path: str) -> str:
    """Extract text from a Word document using python-docx."""
    text = []
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
    except Exception as e:
        print(f"Error parsing DOCX {file_path}: {e}")
        raise e
    return "\n".join(text)

def parse_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"Error parsing TXT {file_path}: {e}")
        raise e

def parse_resume(file_path: str) -> str:
    """Detect file type and extract text content."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext == '.docx':
        return parse_docx(file_path)
    elif ext in ['.txt', '.rtf']:
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
